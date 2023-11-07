# from fastapi import FastAPI, File, UploadFile
# from fastapi.responses import JSONResponse
# from docx import Document
# import openai
# import tempfile

# app = FastAPI()

# # Replace with your actual OpenAI GPT-3 API key
# openai.api_key = "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"

# @app.post("/upload_document/")
# async def upload_document(file: UploadFile, questions: str):
#     try:
#         if not file.filename.endswith(".docx"):
#             return JSONResponse(content={"error": "Invalid file format. Please upload a .docx file."})

#         # Create a temporary file to copy the contents of the uploaded file
#         with tempfile.NamedTemporaryFile(delete=False) as temp_file:
#             temp_filename = temp_file.name
#             for chunk in file.file:
#                 temp_file.write(chunk)

#         # Read the temporary Word document file
#         document = Document(temp_filename)
#         extracted_text = "\n".join([para.text for para in document.paragraphs])
#         print(extracted_text)

#         # Prepare input for question-answering model
#         input_text = f"Context: {extracted_text}\nQuestion: {questions}\nAnswer:"

#         # Use the NLP model to generate answers
#         response = openai.Completion.create(
#             engine="text-davinci-002",  # Replace with your desired model
#             prompt=input_text,
#             max_tokens=4000  # Adjust as needed
#         )

#         answer = response.choices[0].text.strip()
#         return JSONResponse(content={"answer": answer})
#     except Exception as e:
#         return JSONResponse(content={"error": f"Error processing the document: {str(e)}"})

# if __name__ == "__main__":
#     import uvicorn
#     uvicorn.run(app, host="0.0.0.0", port=8000)



from fastapi import FastAPI, File, UploadFile,Form
from fastapi.responses import JSONResponse
from docx import Document
import openai
import tempfile

app = FastAPI()

# # Replace with your actual OpenAI GPT-3 API key
openai.api_key = "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"

@app.post("/process_document/")
async def process_document(file: UploadFile = File(...), questions: str = Form(...)):
    try:
        if not file.filename.endswith(".docx"):
            return JSONResponse(content={"error": "Invalid file format. Please upload a .docx file."})

        # Create a temporary file to copy the contents of the uploaded file
        with tempfile.NamedTemporaryFile(delete=False) as temp_file:
            temp_filename = temp_file.name
            for chunk in file.file:
                temp_file.write(chunk)

        # Initialize an empty string to store the extracted text
        extracted_text = ""

        # Read the temporary Word document file in chunks
        document = Document(temp_filename)
        for para in document.paragraphs:
            extracted_text += para.text + "\n"

        # Split the extracted text into smaller chunks if it's too long
        chunk_size = 4000  # Adjust the size as needed
        chunks = [extracted_text[i:i + chunk_size] for i in range(0, len(extracted_text), chunk_size)]

        # Initialize a list to store the answers
        answers = []

        # Process each chunk with the NLP model
        for chunk in chunks:
            input_text = f"Context: {chunk}\nQuestion: {questions}\nAnswer:"
            response = openai.Completion.create(
                engine="text-davinci-003",
                prompt=input_text,
                max_tokens=100  # Adjust as needed
            )
            answer = response.choices[0].text.strip()
            answers.append(answer)

        # Combine the answers
        combined_answer = " ".join(answers)

        return JSONResponse(content={"answer": combined_answer})
    except Exception as e:
        return JSONResponse(content={"error": f"Error processing the document: {str(e)}"})


if __name__ == "__main__":
     import uvicorn
     uvicorn.run(app, host="0.0.0.0", port=8000)


